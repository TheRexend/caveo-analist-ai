#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o documento Word com a análise e o plano de reestruturação do Google Ads da Caveo."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE = RGBColor(0x0E, 0x8A, 0x5F)
VERMELHO = RGBColor(0xB0, 0x2A, 0x2A)
CINZA = RGBColor(0x55, 0x55, 0x55)
ESCURO = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = ESCURO


def bullet(text, bold_prefix=None, sub=False):
    p = doc.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numitem(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def kv(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = VERDE
    p.add_run(value)
    return p


def para(text):
    return doc.add_paragraph(text)


def callout(text, color=VERMELHO, label="Achado crítico"):
    p = doc.add_paragraph(style="Intense Quote")
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = color
    p.add_run(text).italic = True
    return p


def table(headers, rows, widths=None, total_row=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if total_row:
        cells = t.add_row().cells
        for i, val in enumerate(total_row):
            cells[i].text = str(val)
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    return t


# ============================================================================
# CAPA
# ============================================================================
doc.add_heading("Google Ads Caveo — Análise e Plano de Reestruturação", level=0)
sub = doc.add_paragraph()
r = sub.add_run("Como melhorar a captação de Médicos Recém-formados e Médicos Maduros (plantonistas PJ)")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = CINZA

meta = doc.add_paragraph()
meta.add_run("Data: ").bold = True
meta.add_run("16/06/2026\n")
meta.add_run("Conta analisada: ").bold = True
meta.add_run("Caveo Tecnologia (3921127876) — MCC 5029399396\n")
meta.add_run("Janela de dados: ").bold = True
meta.add_run("últimos 90 dias (18/03 a 15/06/2026)\n")
meta.add_run("Fonte: ").bold = True
meta.add_run("Google Ads API (estrutura, performance e ações de conversão da conta).")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Aviso metodológico: ")
r.bold = True
r.font.color.rgb = VERMELHO
p.add_run(
    "as contagens de \"conversões\" da conta hoje misturam ações de qualidades muito diferentes "
    "(formulários, inscrições no YouTube, visitas à loja, etc. — ver Seção 3.1). Por isso os CPLs "
    "abaixo são DIRECIONAIS, não exatos. A correção da medição é o primeiro passo do plano."
)

doc.add_page_break()

# ============================================================================
# 1. SUMÁRIO EXECUTIVO
# ============================================================================
doc.add_heading("1. Sumário executivo", level=1)
para(
    "A conta está saudável em captação de demanda existente (marca/Institucional), mas tem três "
    "fragilidades estruturais que limitam o crescimento da captação de médicos: (1) a medição que "
    "guia o Smart Bidding está poluída; (2) a captação de médicos maduros depende de Search de "
    "contabilidade, cujo CPL é proibitivo (R$240–R$280); e (3) não há geração de demanda ativa — "
    "todas as campanhas de topo (Demand Gen e YouTube) estão pausadas."
)
doc.add_heading("As 5 maiores oportunidades, por ordem de impacto", level=2)
numitem("limpar as conversões primárias e passar a otimizar por lead qualificado / oportunidade do Salesforce (opportunity_created), não por volume bruto. Maior alavanca da conta e custo zero.", bold_prefix="Consertar a medição — ")
numitem("recolocar topo de funil no ar, com criativo contraintuitivo (já mapeado), Customer Match dos +15.000 médicos como semente de lookalike e otimização por lead. É o caminho mais barato para escalar volume qualificado, sobretudo para médicos maduros (onde o Search é caríssimo).", bold_prefix="Reativar Demand Gen — ")
numitem("hoje o Institucional (marca) consome a maior fatia. Separar marca de genérico evita que o relatório \"pareça\" eficiente às custas de demanda que já viria de graça.", bold_prefix="Separar marca de não-marca — ")
numitem("estruturar a captação em duas linhas claras (RF e MM), com landing pages dedicadas e mensagens distintas, em vez de campanhas soltas competindo por orçamento.", bold_prefix="Reorganizar por público — ")
numitem("corrigir a LP em SPA (renderização JS) que prejudica relevância no Search e migrar o peso de MM do Search caro para Demand Gen/PMax bem alimentado.", bold_prefix="Reduzir dependência do Search de MM — ")

doc.add_page_break()

# ============================================================================
# 2. RETRATO DA CONTA HOJE
# ============================================================================
doc.add_heading("2. Retrato da conta hoje", level=1)
para(
    "Campanhas ativas (ENABLED) e a PMax recém-pausada que ainda concentra histórico relevante, "
    "nos últimos 90 dias. Valores em R$. CPL = custo / conversões (ver aviso metodológico)."
)
table(
    ["Campanha", "Tipo", "Verba/dia", "Custo 90d", "Cliques", "Conv.", "CPL", "CTR", "CPC"],
    [
        ["[Search] Institucional (marca)", "Search", "R$200", "R$12.993", "6.426", "433", "R$30", "30,7%", "R$2,02"],
        ["[RF] Cnpj médico", "Search", "R$90", "R$9.768", "1.025", "98", "R$99,67", "6,7%", "R$9,53"],
        ["[MM] Contabilidade — BR", "Search", "R$40", "R$5.656", "226", "20", "R$282,82", "6,6%", "R$25,03"],
        ["[MM] Contabilidade — SP|RJ|BA", "Search", "R$20", "R$3.901", "162", "16", "R$243,80", "6,1%", "R$24,08"],
        ["[MM] PMax Médico PJ - Novo", "PMax", "R$190", "R$393", "384", "2", "R$196,70", "25,1%", "R$1,02"],
        ["[MM] PMax Plataforma Financeira (PAUSADA)", "PMax", "R$140", "R$3.916", "8.487", "565", "R$6,93", "7,3%", "R$0,46"],
    ],
    total_row=["TOTAL ativas (sem PMax pausada)", "", "", "R$32.711", "8.223", "569", "≈ R$57*", "", ""],
)
p = doc.add_paragraph()
r = p.add_run("* CPL médio ponderado das ativas — puxado para baixo pelo Institucional (marca). ")
r.italic = True
r.font.size = Pt(9)
p.add_run(
    "Investimento ativo ≈ R$363/dia. As campanhas com prefixo \"TP\" (gestão anterior) estão todas "
    "pausadas/removidas — incluindo 4 campanhas Demand Gen, 1 de YouTube reconhecimento e várias "
    "Search de teste — formando um \"cemitério\" de campanhas sem aprendizado preservado."
).font.size = Pt(9)

doc.add_heading("Leitura rápida", level=2)
bullet("Institucional (marca) é o maior gastador e parece excelente (R$30 de CPL, CTR 30,7% típico de busca por marca). Mas captura demanda que já existe — não cria médicos novos.", bold_prefix="Marca domina o gasto. ")
bullet("R$99 de CPL num único ponto de captação. Há espaço para escalar volume e diversificar formatos.", bold_prefix="RF concentrado. ")
bullet("CPL de R$240–R$280 no Search. \"Contabilidade médica\" é um dos termos mais disputados do Brasil (CPC R$24–R$25). Esse caminho não escala de forma saudável.", bold_prefix="MM é caro no Search. ")
bullet("a PMax que gerava volume barato (R$6,93) foi pausada e substituída por uma \"Nova\" que mal gasta (R$393 em 90 dias) e ainda não saiu do aprendizado. O volume barato sumiu sem substituto.", bold_prefix="PMax descontinuada. ")
bullet("nenhuma campanha de Demand Gen ou YouTube ativa. A conta só captura demanda; não gera.", bold_prefix="Topo de funil zerado. ")

doc.add_page_break()

# ============================================================================
# 3. DIAGNÓSTICO
# ============================================================================
doc.add_heading("3. Diagnóstico dos problemas estruturais", level=1)

doc.add_heading("3.1. A medição que guia o Smart Bidding está poluída (prioridade máxima)", level=2)
para(
    "A conta tem como conversões PRIMÁRIAS (as que alimentam o Maximize Conversions) um conjunto "
    "que mistura intenções incomparáveis:"
)
table(
    ["Conversão primária hoje", "Categoria", "Deveria ser primária?"],
    [
        ["Formulário LP", "Lead form", "Sim (lead)"],
        ["[GA4] - form_submit", "Lead form", "Provável duplicidade do form — auditar"],
        ["opportunity_created (Salesforce)", "Lead convertido", "SIM — é o sinal de qualidade"],
        ["Local actions - Other engagements", "Engajamento", "Não — ruído"],
        ["YouTube channel subscriptions", "Engajamento", "Não — ruído"],
        ["YouTube follow-on views", "Engajamento", "Não — ruído"],
        ["Store visits", "Visita à loja", "Não — produto é digital"],
    ],
)
callout(
    "o algoritmo trata uma inscrição no YouTube ou uma visita à loja como equivalente a um lead "
    "real. Pior: o opportunity_created — a importação offline das Oportunidades do Salesforce, que é "
    "o ÚNICO sinal de qualidade verdadeiro — está afogado no meio do ruído. A conta otimiza por "
    "volume de qualquer coisa, não por médico que vira oportunidade."
)
bullet("manter como primárias apenas o lead de formulário (deduplicado) e o opportunity_created; rebaixar YouTube, store visits e local actions para secundárias (observação).", bold_prefix="Correção: ")
bullet("resolver a provável dupla contagem entre \"Formulário LP\", \"Formulário SITE\" e \"[GA4] - form_submit\".", bold_prefix="Correção: ")
bullet("atribuir valor às conversões (valor de oportunidade / ticket médio por público) para migrar de Maximize Conversions para estratégias por valor.", bold_prefix="Correção: ")

doc.add_heading("3.2. Dependência de captura (Search/marca) e ausência de geração de demanda", level=2)
para(
    "O grosso do orçamento está em Search — formato que só colhe quem já procura. Como a maioria "
    "dos médicos (sobretudo recém-formados) NÃO sabe que precisa de PJ/otimização tributária até "
    "alguém apontar, depender de Search limita o teto de captação. Sem Demand Gen/YouTube, não há "
    "quem crie essa demanda."
)

doc.add_heading("3.3. Médicos maduros: Search inviável por CPL", level=2)
para(
    "Os termos de \"contabilidade médica\" custam R$24–R$25 por clique e entregam CPL de R$240–R$280. "
    "Mesmo com landing page perfeita, a matemática do Search puro para MM é frágil. A PMax chegou a "
    "trazer volume a R$6,93, mas (a) provavelmente com qualidade baixa — PMax tende a inflar topo de "
    "funil — e (b) foi pausada. Falta um canal de volume qualificado para MM."
)

doc.add_heading("3.4. Recém-formados concentrados num único ponto de captação", level=2)
para(
    "Há uma só campanha de RF (Search \"Cnpj médico\", CPL ~R$99). É um público com forte "
    "sazonalidade (formaturas em janeiro/julho, provas de residência) e altíssima resposta a "
    "vídeo/educação — perfeito para Demand Gen e YouTube, que hoje não existem para ele."
)

doc.add_heading("3.5. A landing page em SPA prejudica a relevância no Search", level=2)
para(
    "A lp.caveo.com.br é uma SPA renderizada por JavaScript: para os robôs do Google, o HTML chega "
    "praticamente vazio. Isso é causa-raiz conhecida de baixo Índice de Qualidade / \"relevância "
    "pouco clara\", o que encarece o CPC justamente nas campanhas de Search mais caras (MM). "
    "Corrigir a renderização (SSR/pré-render) ou usar LPs dedicadas server-side derruba CPC e CPL."
)

doc.add_page_break()

# ============================================================================
# 4. ESTRATÉGIA POR PÚBLICO
# ============================================================================
doc.add_heading("4. Estratégia por público", level=1)

doc.add_heading("4.1. Médicos recém-formados (e último ano / residentes)", level=2)
kv("Jornada", "transição com baixa consciência do problema fiscal; decide rápido perto da formatura; consome muito vídeo/Reels/Shorts; sensível a \"sair na frente\" e a \"começar certo\".")
kv("Diagnóstico atual", "só Search de marca/cnpj; sem topo de funil; sazonalidade não explorada.")
para("Recomendações:")
bullet("manter o Search non-brand enxuto (cnpj médico, abrir pj medicina, mei vs pj médico) com landing dedicada e extensões de \"CNPJ grátis\".", bold_prefix="Captura: ")
bullet("Demand Gen com criativos de vídeo de educação/urgência (conceitos \"Parabéns, o governo agradece\" e \"A IA te dá 12 passos\"), otimizada por lead, semeada com Customer Match + lookalike de leads RF.", bold_prefix="Geração: ")
bullet("intensificar verba em jan e jul (formaturas) e nas janelas de resultado de residência; reduzir no meio do ciclo.", bold_prefix="Sazonalidade: ")

doc.add_heading("4.2. Médicos maduros (plantonistas PJ estabelecidos)", level=2)
kv("Jornada", "já sentem a dor (imposto alto, contador genérico) mas são céticos e inertes (medo de trocar). Decisão mais longa e de maior valor (LTV alto).")
kv("Diagnóstico atual", "Search caríssimo (R$240–280 CPL); PMax de volume pausada; sem retargeting estruturado.")
para("Recomendações:")
bullet("manter Search apenas nos termos de altíssima intenção (\"contabilidade para médicos\", \"contador plantonista\"), com lances por valor e foco em oportunidade — não em volume de formulário.", bold_prefix="Captura cirúrgica: ")
bullet("transferir a busca por volume do Search para Demand Gen + PMax BEM alimentada (sinais de público e otimização por opportunity_created), reduzindo o custo por oportunidade.", bold_prefix="Volume via topo: ")
bullet("Demand Gen/Display com a \"conta de guardanapo\" (até 30%) e \"Caveo vs. contador genérico\"; retargeting de quem visitou a LP e não converteu, e quem abandonou o cadastro.", bold_prefix="Retargeting: ")

doc.add_page_break()

# ============================================================================
# 5. ARQUITETURA DE CAMPANHAS PROPOSTA
# ============================================================================
doc.add_heading("5. Arquitetura de campanhas proposta", level=1)
para(
    "Estrutura enxuta, organizada por público × estágio de funil, com nomenclatura consistente. "
    "Cada campanha tem um papel claro e não compete às cegas por orçamento."
)
table(
    ["Camada", "Campanha", "Tipo", "Otimização", "Papel"],
    [
        ["Marca", "[Brand] Caveo", "Search", "Lead / Opp", "Proteger marca; teto de gasto baixo; isolado do não-marca"],
        ["Captura RF", "[RF] Cnpj/PJ médico", "Search", "Lead → Opp", "Alta intenção recém-formado"],
        ["Captura MM", "[MM] Contabilidade médica", "Search", "Valor / Opp", "Alta intenção, termos selecionados"],
        ["Geração RF", "[RF] Demand Gen", "Demand Gen", "Lead", "Urgência/educação; sazonal formatura"],
        ["Geração MM", "[MM] Demand Gen", "Demand Gen", "Lead / Opp", "Volume qualificado fora do Search caro"],
        ["Escala", "[ALL] PMax alimentada", "PMax", "Opp + valor", "Escala com sinais de público e feed; reativar com cautela"],
        ["Retorno", "[ALL] Retargeting", "Demand Gen/Display", "Lead", "Quem visitou LP / abandonou cadastro"],
        ["Topo (fase 2)", "[ALL] YouTube", "Vídeo", "Consideração", "Reconhecimento + alimenta retargeting"],
    ],
)
bullet("separar marca de não-marca para enxergar o CPL real da captação (hoje o Institucional mascara a conta).")
bullet("uma campanha por público/estágio — evita canibalização e facilita decisão de verba.")
bullet("toda campanha de volume otimizada com opportunity_created como sinal (ou valor), nunca só por formulário.")

doc.add_page_break()

# ============================================================================
# 6. DEMAND GEN — PLAYBOOK
# ============================================================================
doc.add_heading("6. Demand Gen — playbook detalhado", level=1)
para(
    "Demand Gen é o canal certo para o gargalo da Caveo: leva criativo nativo (vídeo e imagem) a "
    "YouTube, Shorts, Discover e Gmail, com segmentação por públicos semelhantes — ideal para CRIAR "
    "demanda em quem ainda não busca. A Caveo já tentou Demand Gen no passado (4 campanhas, todas "
    "pausadas); a diferença desta vez é fazê-lo com medição correta, públicos próprios e criativo forte."
)

doc.add_heading("6.1. Por que agora vai funcionar (o que faltou antes)", level=2)
bullet("otimizar por opportunity_created/valor, não por engajamento de vídeo (que era primário e poluía tudo).", bold_prefix="Medição: ")
bullet("Customer Match com a base de +15.000 médicos como semente de lookalike — sinal que campanhas anteriores não tinham.", bold_prefix="Sinal próprio: ")
bullet("usar os conceitos contraintuitivos já validados (print de chat com IA, anamnese fiscal, conta de guardanapo) em vez de criativo genérico.", bold_prefix="Criativo: ")

doc.add_heading("6.2. Segmentação", level=2)
bullet("Customer Match: lista de leads/clientes médicos (semente de lookalike) e lista de CLIENTES atuais (para EXCLUIR e evitar gastar com quem já é cliente).", bold_prefix="Públicos próprios: ")
bullet("públicos semelhantes às listas; segmentos de afinidade/in-market de saúde, finanças e empreendedorismo; segmentação demográfica por ocupação onde disponível.", bold_prefix="Lookalike + interesses: ")
bullet("RF: 22–32 anos, interesses de medicina/residência/concursos. MM: 30–50 anos, plantonista/consultório/sociedade.", bold_prefix="Por público: ")

doc.add_heading("6.3. Criativos (reaproveitar o documento de ideias)", level=2)
bullet("vídeos verticais 9:16 de 15–30s — \"Parabéns, o governo agradece\" e \"A IA te dá 12 passos, a Caveo 1 botão\". CTA: abrir PJ grátis.", bold_prefix="RF: ")
bullet("\"conta de guardanapo\" (até 30%), \"diagnóstico sem exame\" e \"Caveo vs. contador genérico\" em vídeo e estático. CTA: simular economia / falar com especialista.", bold_prefix="MM: ")
bullet("3–5 criativos por campanha, renovados a cada 2–3 semanas para combater fadiga; testar 1 ângulo por vez.", bold_prefix="Volume e rotação: ")

doc.add_heading("6.4. Configuração e lances", level=2)
bullet("começar em Maximize Conversions otimizando por LEAD; ao acumular ~30–50 oportunidades, migrar para tCPA por oportunidade (ou por valor).", bold_prefix="Lances: ")
bullet("orçamento inicial sugerido de teste: R$100–R$150/dia por público (RF e MM separados), com janela mínima de 2–3 semanas antes de julgar.", bold_prefix="Verba: ")
bullet("landing pages dedicadas por público (não a home SPA), com formulário curto e prova social (+15.000 médicos / 21 estados).", bold_prefix="Destino: ")

doc.add_page_break()

# ============================================================================
# 7. ALÉM DE DEMAND GEN
# ============================================================================
doc.add_heading("7. Além de Demand Gen — outras alavancas", level=1)

doc.add_heading("7.1. Lances por valor e por oportunidade (o multiplicador)", level=2)
para(
    "Com a medição limpa e valores atribuídos, migrar as campanhas de volume de \"Maximize "
    "Conversions\" para tCPA por oportunidade ou Maximize Conversion Value / tROAS. Isso faz o Google "
    "buscar o médico que VIRA cliente, não o clique mais barato. É o que transforma CPL em CPO "
    "(custo por oportunidade) como métrica-mãe."
)

doc.add_heading("7.2. Customer Match a partir dos +15.000 médicos", level=2)
bullet("semear lookalikes de quem já é cliente (perfil comprovadamente bom).", bold_prefix="Aquisição: ")
bullet("excluir clientes atuais de todas as campanhas de aquisição (parar de pagar por quem já tem conta).", bold_prefix="Eficiência: ")
bullet("listas separadas de RF e MM para personalizar mensagem e lances.", bold_prefix="Segmentação: ")

doc.add_heading("7.3. Higiene de Search", level=2)
bullet("auditar termos de pesquisa e negativar desperdício (concursos, CLT, vagas, conteúdo gratuito).", bold_prefix="Negativas: ")
bullet("separar marca de genérico (já citado) e revisar correspondências para evitar termos amplos caros sem intenção.", bold_prefix="Estrutura: ")
bullet("revisar extensões (sitelinks de \"CNPJ grátis\", \"economize até 30%\", \"feito por médicos\") e Índice de Qualidade pós-correção da LP.", bold_prefix="Anúncios: ")

doc.add_heading("7.4. Landing pages e a SPA", level=2)
bullet("corrigir a renderização da LP (SSR/pré-render) para recuperar relevância e baratear o Search; ou usar LPs dedicadas server-side por público.", bold_prefix="Técnico: ")
bullet("uma LP para RF (PJ grátis / começar certo) e uma para MM (até 30% / diagnóstico fiscal), cada uma alinhada ao criativo de origem.", bold_prefix="Conversão: ")

doc.add_heading("7.5. Retargeting e funil de retorno", level=2)
bullet("Demand Gen/Display para quem visitou LP e não converteu e para quem abandonou o cadastro; mensagens de quebra de objeção (migração simples, suporte humano especializado).")

doc.add_heading("7.6. YouTube como topo (fase 2)", level=2)
bullet("após Demand Gen estabilizar, somar YouTube de consideração para encher os públicos de retargeting e reduzir o CPO ao longo do tempo — com os mesmos conceitos criativos.")

doc.add_heading("7.7. Geografia e sazonalidade", level=2)
bullet("concentrar verba nos estados de maior densidade médica e melhor histórico de oportunidade; revisar o split SP|RJ|BA por CPO real, não por CPL.", bold_prefix="Geo: ")
bullet("calendário de formaturas (jan/jul) e provas de residência para escalar RF nas janelas certas.", bold_prefix="Sazonalidade: ")

doc.add_page_break()

# ============================================================================
# 8. ROADMAP E KPIs
# ============================================================================
doc.add_heading("8. Roadmap 30 / 60 / 90 dias", level=1)
table(
    ["Janela", "Ações", "Resultado esperado"],
    [
        ["0–30 dias\n(fundação)",
         "Limpar conversões primárias; resolver dupla contagem; atribuir valor; separar marca de não-marca; auditar negativas; subir Customer Match (semente + exclusão).",
         "Smart Bidding passa a otimizar por lead/oportunidade real; CPL real visível."],
        ["30–60 dias\n(geração)",
         "Lançar Demand Gen RF e MM com criativos contraintuitivos e LPs dedicadas; corrigir/contornar a SPA; iniciar retargeting.",
         "Volume qualificado fora do Search caro; queda do CPO de MM."],
        ["60–90 dias\n(escala)",
         "Migrar campanhas de volume para tCPA/tROAS por oportunidade; reativar PMax bem alimentada; somar YouTube de topo; escalar o que prova CPO.",
         "Crescimento de oportunidades a custo controlado; menor dependência de marca."],
    ],
)

doc.add_heading("8.1. KPIs — medir o que importa", level=2)
bullet("custo por oportunidade do Salesforce (não por formulário). É a métrica-mãe.", bold_prefix="CPO: ")
bullet("taxa lead → oportunidade por campanha/UTM (qualidade da fonte).", bold_prefix="Lead→Opp: ")
bullet("% do gasto em não-marca vs. marca (medir captação real).", bold_prefix="Share non-brand: ")
bullet("CPO por público (RF e MM separados) e por canal (Search vs. Demand Gen vs. PMax).", bold_prefix="CPO por público: ")
para(
    "Acompanhar tudo isso no dashboard Caveo (modelo de duas datas no funil), cruzando UTMs do "
    "Google com os estágios do Salesforce — exatamente o trabalho do agente analista-midia-paga-crm."
)

doc.add_heading("9. Riscos e dependências", level=1)
bullet("a correção da medição depende de acesso ao GA4/Tag e à importação offline do Salesforce (opportunity_created).", bold_prefix="Medição: ")
bullet("a correção da SPA depende do time de produto/engenharia; enquanto isso, usar LPs dedicadas server-side.", bold_prefix="LP: ")
bullet("Demand Gen e mudanças de lance entram em aprendizado por 1–2 semanas; não julgar resultados antes disso.", bold_prefix="Aprendizado: ")
bullet("PMax precisa de sinais e exclusões corretas para não inflar topo de funil; reativar só após a medição estar limpa.", bold_prefix="PMax: ")

doc.save("/Users/matheus/Documents/Claude/Projects/caveo_analist_ai/docs/Google_Ads_Caveo_Analise_Reestruturacao.docx")
print("OK: docs/Google_Ads_Caveo_Analise_Reestruturacao.docx")
