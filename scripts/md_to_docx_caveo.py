#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Conversor Markdown -> DOCX com identidade visual Caveo/Boomer.

Uso:
    python3 md_to_docx_caveo.py <entrada.md> [saida.docx]

Cobre os constructos usados nos docs do projeto: títulos (#, ##, ###),
blockquotes (>), listas com bullets/numeradas, checklists (- [ ]),
tabelas pipe, e ênfase inline (**negrito**, *itálico*, `código`,
[[wikilink]], [texto](url)). Quebra de página antes de cada seção de
nível 1 (##), no padrão dos demais documentos da conta.
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

VERDE = RGBColor(0x0E, 0x8A, 0x5F)
CINZA = RGBColor(0x55, 0x55, 0x55)
ESCURO = RGBColor(0x1A, 0x1A, 0x1A)

# Tokeniza ênfase inline preservando a ordem do texto.
TOKEN = re.compile(
    r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[\[[^\]]+?\]\]|\[[^\]]+?\]\([^)]+?\))'
)


def add_inline(paragraph, text):
    """Adiciona `text` ao parágrafo interpretando ênfase markdown inline."""
    for tok in TOKEN.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        elif tok.startswith('[['):
            paragraph.add_run(tok[2:-2])  # wikilink -> texto puro
        elif tok.startswith('[') and '](' in tok:
            paragraph.add_run(tok[1:tok.index(']')])  # [label](url) -> label
        else:
            paragraph.add_run(tok)


def is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|')


def is_separator_row(line):
    return bool(re.fullmatch(r'\|[\s:|-]+\|', line.strip()))


def cells_of(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as fh:
        lines = fh.read().splitlines()

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = ESCURO

    seen_section = False  # controla page-break entre seções de nível 1
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- linha em branco
        if not stripped:
            i += 1
            continue

        # --- regra horizontal: vira espaçador discreto
        if re.fullmatch(r'-{3,}', stripped) or re.fullmatch(r'\*{3,}', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # --- títulos
        m = re.match(r'(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                doc.add_heading(text, level=0)
            elif level == 2:
                if seen_section:
                    doc.add_page_break()
                seen_section = True
                h = doc.add_heading('', level=1)
                add_inline(h, text)
            else:
                h = doc.add_heading('', level=min(level - 1, 4))
                add_inline(h, text)
            i += 1
            continue

        # --- tabela
        if is_table_row(line) and i + 1 < n and is_separator_row(lines[i + 1]):
            headers = cells_of(line)
            rows = []
            j = i + 2
            while j < n and is_table_row(lines[j]):
                rows.append(cells_of(lines[j]))
                j += 1
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = 'Light Grid Accent 1'
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for c, htext in enumerate(headers):
                cell = t.rows[0].cells[c]
                cell.paragraphs[0].text = ''
                add_inline(cell.paragraphs[0], htext)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c in range(len(headers)):
                    val = row[c] if c < len(row) else ''
                    cells[c].paragraphs[0].text = ''
                    add_inline(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            i = j
            continue

        # --- blockquote
        if stripped.startswith('>'):
            content = re.sub(r'^>\s?', '', stripped)
            p = doc.add_paragraph(style='Quote')
            add_inline(p, content)
            for run in p.runs:
                run.font.color.rgb = CINZA
            i += 1
            continue

        # --- checklist
        mc = re.match(r'-\s+\[([ xX])\]\s+(.*)', stripped)
        if mc:
            box = '☑ ' if mc.group(1).lower() == 'x' else '☐ '
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(box)
            add_inline(p, mc.group(2))
            i += 1
            continue

        # --- lista com bullets
        mb = re.match(r'[-*]\s+(.*)', stripped)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mb.group(1))
            i += 1
            continue

        # --- lista numerada
        mn = re.match(r'\d+\.\s+(.*)', stripped)
        if mn:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mn.group(1))
            i += 1
            continue

        # --- parágrafo comum
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'OK: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        sys.exit('uso: python3 md_to_docx_caveo.py <entrada.md> [saida.docx]')
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else re.sub(r'\.md$', '.docx', src)
    convert(src, dst)
