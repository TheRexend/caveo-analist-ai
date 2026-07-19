#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o documento Word com todas as ideias de criativos de anúncios para a Caveo."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE = RGBColor(0x0E, 0x8A, 0x5F)
CINZA = RGBColor(0x55, 0x55, 0x55)
ESCURO = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = ESCURO


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_hook(text):
    """Hook destacado em itálico com barra lateral verde (citação)."""
    p = doc.add_paragraph(style="Intense Quote")
    run = p.add_run(text)
    run.italic = True
    return p


def add_kv(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = VERDE
    p.add_run(value)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def conceito_heading(num, titulo):
    h = doc.add_heading(f"Conceito {num} — {titulo}", level=2)


# ----------------------------------------------------------------------------
# CAPA
# ----------------------------------------------------------------------------
title = doc.add_heading("Ideias de Criativos de Anúncios — Caveo", level=0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run("Levantamento de dores, desejos e conceitos contraintuitivos para Médicos Recém-formados e Médicos Maduros")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = CINZA

meta = doc.add_paragraph()
meta.add_run("Data: ").bold = True
meta.add_run("15/06/2026\n")
meta.add_run("Produto: ").bold = True
meta.add_run(
    "Plataforma financeira/contábil feita por médicos para médicos — abertura de PJ, emissão "
    "ilimitada de NF, otimização tributária (até 30%), gestão de plantões e múltiplas fontes de "
    "renda, dashboard em tempo real. App que transforma ~600h/ano de burocracia em 15 min/mês.\n"
)
meta.add_run("Tração: ").bold = True
meta.add_run("+15.000 médicos em 21 estados.")

# Caixa de regras de integridade
doc.add_heading("Regras de integridade (ler antes de produzir)", level=2)
bullet("a Caveo NÃO tem IA no sistema (apenas o atendimento primário do suporte). Nos criativos, IA é apenas referência cultural — nenhum hook pode alegar funcionalidade de IA no app.", bold_prefix="IA: ")
bullet('a economia é sempre "até 30%" e usar "pode" — provocação, nunca promessa garantida.', bold_prefix="Economia: ")
bullet('a "conta de guardanapo" é sempre ilustrativa/teórica; o cálculo real é feito com os dados do médico.', bold_prefix="Cálculos: ")

doc.add_page_break()

# ----------------------------------------------------------------------------
# 1. CONTEXTO ESTRATÉGICO
# ----------------------------------------------------------------------------
doc.add_heading("1. Contexto estratégico", level=1)

doc.add_heading("Públicos e nível de consciência", level=2)
bullet(
    "estão em transição (medo + esperança). Em geral NÃO sabem que o problema existe até ser "
    "apontado. O anúncio precisa educar e criar urgência.",
    bold_prefix="Recém-formados (último ano + residentes): ",
)
bullet(
    "já sentem a dor (imposto alto, contador genérico), mas são céticos. O anúncio precisa "
    "provar a diferença e quebrar objeção.",
    bold_prefix="Médicos maduros (PJ estabelecido / plantonista / consultório): ",
)

doc.add_heading("O que torna um criativo contraintuitivo aqui", level=2)
p = doc.add_paragraph(
    "A categoria de contabilidade/fintech vive de confiança, seriedade, terno, calculadora e "
    '"economize 30%". Todos os concorrentes falam igual. O contraintuitivo de verdade é quebrar '
    "o padrão da categoria — o criativo deve parecer qualquer coisa, menos um anúncio de "
    "contabilidade. Duas alavancas:"
)
bullet("o criativo não parece anúncio (parece print de conversa, exame, receita médica).", bold_prefix="Pattern interrupt: ")
bullet("dizer em voz alta o que o médico pensa mas ninguém na categoria fala.", bold_prefix="Verdade incômoda: ")

doc.add_page_break()

# ----------------------------------------------------------------------------
# 2. MAPA DE DORES E DESEJOS
# ----------------------------------------------------------------------------
doc.add_heading("2. Mapa de dores e desejos", level=1)


def tabela_dor(titulo, linhas):
    doc.add_heading(titulo, level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    hdr[0].text = titulo.split(" — ")[-1]
    hdr[1].text = "Gatilho emocional"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    for col, gat in linhas:
        row = t.add_row().cells
        row[0].text = col
        row[1].text = gat
    doc.add_paragraph()


doc.add_heading("Público 1 — Recém-formados", level=2)
tabela_dor("Dores", [
    ("Medo de perder as melhores vagas/plantões por não ter CNPJ pronto na formatura", "Ansiedade competitiva"),
    ("Burocracia de abrir PJ: não sabe por onde começar, custo de +R$1.000", "Paralisia, despreparo"),
    ("Insegurança fiscal: medo de errar imposto, cair na malha, multa no início", "Medo de estragar o começo"),
    ("Sem tempo (residência/plantão intenso) para papelada", "Exaustão"),
    ("Ninguém ensinou isso na faculdade — formou médico, não empreendedor", "Insegurança, vergonha"),
    ("Renda bagunçada (bolsa + plantões + múltiplas fontes)", "Falta de controle"),
])
tabela_dor("Desejos", [
    ("Começar a faturar no dia seguinte à formatura, sem fricção", "Aspiração, liberdade"),
    ("Sair na frente da concorrência — PJ pronta antes dos colegas", "Status, vantagem"),
    ("Liberdade e controle financeiro desde o dia 1", "Autonomia"),
    ("Economia imediata (CNPJ grátis, +R$1.000)", "Ganho concreto"),
    ("Ser guiado por quem entende medicina (feito por médicos)", "Confiança, pertencimento"),
    ("Simplicidade: resolver tudo no app, em minutos", "Alívio"),
])

doc.add_heading("Público 2 — Médicos maduros", level=2)
tabela_dor("Dores", [
    ("Carga tributária alta — sensação de pagar imposto demais todo mês", "Indignação, perda"),
    ("Contabilidade tradicional genérica: lenta, burocrática, sem otimização", "Frustração, descaso"),
    ("Falta de clareza dos números; múltiplas fontes desorganizadas", "Insegurança, perda de controle"),
    ("Tempo perdido com papelada (~600h/ano)", "Desperdício"),
    ("Medo/preguiça de trocar de contador", "Inércia, risco percebido"),
])
tabela_dor("Desejos", [
    ("Economizar até 30% de imposto — dinheiro de volta no bolso", "Ganho concreto (desejo nº1)"),
    ("Clareza em tempo real dos ganhos, despesas e economia", "Controle, tranquilidade"),
    ("Especialista que entende plantão/PJ médico", "Confiança, ser compreendido"),
    ("Centralizar tudo (plantões + consultório + sociedade)", "Organização"),
    ("Agilidade — emitir NF e pagar tributo em segundos", "Alívio, recuperar tempo"),
])

doc.add_page_break()

# ----------------------------------------------------------------------------
# 3. ÂNGULOS + HOOKS POR DOR
# ----------------------------------------------------------------------------
doc.add_heading("3. Ângulos + hooks por dor", level=1)


def tabela_angulo(titulo, linhas):
    doc.add_heading(titulo, level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Dor/Desejo"
    hdr[1].text = "Ângulo"
    hdr[2].text = "Hooks para testar"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    for a, b, c in linhas:
        row = t.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    doc.add_paragraph()


tabela_angulo("Recém-formados", [
    ("Perder as melhores vagas", "Urgência competitiva: PJ pronta = largada na frente",
     '"Os melhores plantões já têm dono na semana da formatura. Você vai chegar com CNPJ pronto ou correndo atrás?"\n"Seu diploma sai em [mês]. Sua PJ devia estar pronta antes dele."'),
    ("Burocracia de abrir PJ", "Remover fricção: de R$1.000 e dor de cabeça → grátis e simples",
     '"Abrir CNPJ médico custa +R$1.000 e semanas de papelada. Ou 0 e alguns toques no app."\n"Você estudou 6 anos pra salvar vidas, não pra entender Junta Comercial."'),
    ("Insegurança fiscal", "Proteção: não estrague o começo com erro de imposto",
     '"Errar imposto no 1º ano de carreira é mais comum (e mais caro) do que você imagina."\n"Cair na malha logo no começo? Não com quem só cuida de médico."'),
    ("Ninguém ensinou isso", "Mentor que entende: a faculdade formou médico, não empreendedor",
     '"Na faculdade ninguém te ensinou a virar PJ. A gente ensina — e faz por você."\n"Feito por médicos, pra você não aprender finanças no erro."'),
    ("Faturar já na formatura", "Aspiração: do canudo ao 1º plantão pago sem espera",
     '"Da formatura ao primeiro plantão pago, sem perder uma semana."\n"Forme-se sexta. Comece a faturar como PJ na segunda."'),
    ("Sair na frente", "Vantagem de quem se antecipa (último ano/residente)",
     '"Enquanto seus colegas vão abrir a PJ depois de formar, a sua já está rodando."\n"Residência é maratona. Largue na frente na parte financeira."'),
    ("Economia imediata", "Oferta concreta: CNPJ 100% grátis",
     '"Abertura de CNPJ: +R$1.000 no mercado. R$0 na Caveo."'),
])

tabela_angulo("Médicos maduros", [
    ("Economizar até 30%", "O grande gancho: dinheiro que vaza todo mês",
     '"Você pode estar pagando até 30% de imposto a mais. Todo mês."\n"Quanto desse seu plantão vai embora em imposto que dava pra economizar?"'),
    ("Contador genérico", "Comparação: especialista vs. contabilidade tradicional",
     '"Seu contador entende de mercado, padaria e... medicina? A gente só entende de médico."\n"Contador genérico te faz pagar imposto. Contador de médico te faz economizar."'),
    ("Carga tributária alta", "Indignação + solução: imposto de médico é otimizável",
     '"Plantão, consultório, sociedade — cada fonte mal organizada é imposto pago à toa."\n"Não é sobre ganhar mais. É sobre parar de doar pro Leão."'),
    ("Falta de clareza", "Controle: dashboard em tempo real",
     '"Você sabe, agora, quanto ganhou e quanto economizou de imposto esse mês? Devia saber."\n"Seus números num lugar só — não em 3 planilhas e um WhatsApp do contador."'),
    ("Tempo perdido (600h/ano)", "Recuperar tempo: 600h/ano → 15 min/mês",
     '"Médicos gastam ~600 horas por ano com burocracia. A gente devolve 599."\n"15 minutos por mês. O resto do tempo, seja médico."'),
    ("Medo de trocar de contador", "Quebra de objeção: migração simples e sem risco",
     '"Trocar de contador parece um perrengue. Com a gente, é um toque — a gente cuida da migração."'),
    ("Especialista que entende", "Pertencimento + prova social",
     '"Mais de 15.000 médicos em 21 estados já trocaram a contabilidade genérica pela Caveo."'),
])

doc.add_page_break()

# ----------------------------------------------------------------------------
# 4. CONCEITOS CONTRAINTUITIVOS
# ----------------------------------------------------------------------------
doc.add_heading("4. Conceitos contraintuitivos completos", level=1)

conceito_heading(1, '"A IA trava. A gente não."')
add_kv("Público", "Ambos (2 versões)")
add_kv("Formato", "Estático que parece print de chat (ou vídeo da tela sendo digitada)")
doc.add_heading("Versão A — Maduro", level=3)
add_kv("Visual", 'Screenshot estilo ChatGPT/Claude. Usuário: "Sou médico, faço plantão em 3 hospitais + consultório. Quanto imposto pago a mais por não otimizar?" → IA responde com disclaimer genérico → balão verde da Caveo por cima: "A gente responde. Com seus números, em 15 min, com quem só cuida de médico PJ."')
add_hook('Hook: "Já perguntou pra IA quanto imposto você paga a mais? Ela trava na hora. A gente não."')
add_kv("CTA", "Falar com especialista")
doc.add_heading("Versão B — Recém-formado", level=3)
add_kv("Visual", "Print de chat: a IA responde com 12 passos burocráticos para abrir CNPJ (assustador). Caveo entra: \"Ou… a gente abre sua PJ pra você. De graça. Antes de você terminar de ler a resposta da IA.\"")
add_hook('Hook: "A IA te dá 12 passos pra abrir CNPJ. A Caveo te dá 1 botão."')
add_kv("CTA", "Abrir minha PJ grátis")

conceito_heading(2, '"Esse anúncio não é pra você" (anti-pitch)')
add_kv("Público", "Maduro (principal)")
add_kv("Formato", "Vídeo talking-head (fundador médico) ou estático texto-forte")
add_hook('Headline: "Esse anúncio não é pra você. (provavelmente.)"')
add_kv("Corpo", "Se você adora pagar imposto, acha ótimo mandar nota por WhatsApp e confia que seu contador genérico entende de plantão — pode fechar, sério. Mas se você desconfia que paga demais e nunca teve clareza real dos seus números… talvez a gente precise conversar.")
add_kv("CTA", "Ok, me convenceu.")
add_kv("Variante (recém-formado)", "Esse anúncio não é pra estudante de medicina. A não ser que você esteja no último ano e não queira começar a carreira correndo atrás de CNPJ.")

conceito_heading(3, '"Parabéns. O governo também agradece." (timing da formatura)')
add_kv("Público", "Recém-formado")
add_kv("Formato", "Vídeo (cena de formatura) ou estático estilo cartão de parabéns")
add_hook('Hook: "Parabéns pela formatura! 🎉 O governo também quer comemorar."')
add_kv("Corpo", "Seis anos, plantões intermináveis e finalmente o canudo. Aí vem o primeiro plantão pago como PJ — e você descobre que ninguém te ensinou nada sobre CNPJ, nota fiscal ou imposto. A Caveo abre sua PJ de graça e cuida de tudo, pra sua única preocupação ser ser médico.")
add_kv("CTA", "Comece certo: abra sua PJ grátis.")
add_kv("Visual", "A ironia é o anúncio: cartão de parabéns que se transforma em boleto; ou capelo de formatura com etiqueta de imposto pendurada.")
add_kv("Variante", "Você passou em medicina. Ninguém avisou que ia virar microempresário.")

conceito_heading(4, '"Conta de guardanapo"')
add_kv("Público", "Maduro (principal)")
add_kv("Formato", "Estático (guardanapo rabiscado à mão) ou vídeo (mão rabiscando num café/copa do hospital). Estética anti-fintech-polido — feito à caneta.")
add_kv("Visual", 'Guardanapo (ou folha de receituário) com conta à mão: Plantão ~ R$ 25.000/mês → Imposto hoje (chute) → Otimizado (−30%) → Diferença no ano: R$ ZZ.ZZZ 🤯. Nota: "*conta de guardanapo, só pra te assustar. A real, a Caveo faz com os SEUS números."')
add_hook('Hook: "Fiz uma conta de guardanapo do quanto você pode estar pagando de imposto a mais. Senta antes de ver."')
add_kv("Corpo", "Não é proposta, é provocação. Médico PJ que não otimiza pode pagar até 30% a mais em tributos. Quanto dá no seu caso? A Caveo faz a conta de verdade — com seus números — de graça.")
add_kv("CTA", "Quero a conta real.")
add_kv("Variante (recém-formado)", 'Guardanapo comparando "abrir PJ no mercado: +R$1.000" vs. "na Caveo: R$0".')

doc.add_heading("Combo a testar", level=3)
add_hook('Conceito 4 + 1: "Nem o ChatGPT faz essa conta direito, porque ele não sabe que você é médico plantonista." (junta o pattern interrupt do print com a provocação numérica)')

doc.add_page_break()

# ----------------------------------------------------------------------------
# 5. ÂNGULO "CONTEXTO" + SISTEMA DE CAMPANHA
# ----------------------------------------------------------------------------
doc.add_heading('5. Ângulo "Contexto" — o moat real contra a IA', level=1)

p = doc.add_paragraph(
    "Insight-chave: a IA dá resposta genérica não (só) por prompt ruim, mas por FALTA DE CONTEXTO. "
    "Ela não tem acesso aos dados fiscais da pessoa — plantões, consultório, sociedade, notas, "
    "regime. Mesmo com o prompt perfeito, responde no escuro. A Caveo centraliza todas as fontes "
    "de renda num lugar só: ela É o contexto que a IA não tem. Diferenciação real e contraintuitiva "
    "ao mesmo tempo."
)
p2 = doc.add_paragraph()
r = p2.add_run("Paralelo médico perfeito: ")
r.bold = True
p2.add_run(
    "nenhum médico fecha diagnóstico sem anamnese e exames. Mas pergunta para a IA sobre o próprio "
    "imposto com zero \"exames financeiros\". A Caveo é a anamnese fiscal completa."
)

doc.add_heading('Conceito A — "Diagnóstico sem exame"', level=2)
add_kv("Público", "Maduro (principal)")
add_kv("Formato", "Estático split-screen, carrossel ou vídeo 15s")
add_kv("Visual", "Tela dividida. Esquerda: celular com chat no escuro (\"Quanto imposto eu devia pagar? / IA: depende de vários fatores que eu não tenho acesso…\"). Direita: médico analisando exame de imagem na luz. A simetria entre as cenas É o anúncio.")
add_hook('Hook: "Você daria um diagnóstico sem ver os exames? Então por que deixa a IA resolver seu imposto sem ver seus números?"')
add_kv("Corpo", "IA genérica = diagnóstico no escuro. A Caveo enxerga o quadro completo — plantões, consultório, sociedade, notas, regime — e é por isso que economiza até 30% onde a IA só chuta.")
add_kv("CTA", "Faça o diagnóstico fiscal certo.")
bullet('"A IA não erra por burrice. Erra por falta de exame — ela não tem os seus números."', bold_prefix="Variante: ")
bullet('"Imposto no escuro é igual diagnóstico no escuro: perigoso e caro."', bold_prefix="Variante: ")

doc.add_heading('Conceito E — "Anamnese fiscal"', level=2)
add_kv("Público", "Ambos")
add_kv("Formato", "Estático estilo ficha de anamnese (ou vídeo do app preenchendo a ficha sozinho)")
add_kv("Visual", 'Uma ficha de anamnese, mas financeira — campos "Fontes de renda?", "Plantões/mês?", "Sociedade?", "Regime tributário?" — sendo preenchidos automaticamente pelo app, em vez de à caneta.')
add_hook('Hook: "Nenhum médico fecha diagnóstico sem anamnese. Sua vida fiscal merece a mesma coisa."')
add_kv("Corpo", "A Caveo faz a anamnese completa das suas finanças — todas as fontes de renda num lugar só — e monta o plano que reduz até 30% do seu imposto. Contador genérico e IA trabalham sem ficha nenhuma. A gente, não.")
add_kv("CTA", "Comece sua anamnese fiscal.")
bullet('"Anamnese completa antes de qualquer conduta. Inclusive a fiscal."', bold_prefix="Variante: ")
bullet('"Você pergunta tudo pro paciente antes de tratar. Seu contador te pergunta o quê?"', bold_prefix="Variante: ")

doc.add_heading("Sistema de campanha: cada benefício, duas lentes", level=2)
p = doc.add_paragraph(
    'Conceito-guarda-chuva: "a IA chuta porque não te conhece; a Caveo conhece" → uma execução por '
    "benefício. Coluna do meio = enquadramento clínico (linguagem nativa); coluna da direita = "
    "enquadramento contexto/IA (o moat)."
)

t = doc.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "Benefício Caveo"
hdr[1].text = "Ângulo clínico"
hdr[2].text = "Ângulo contexto / IA"
for c in hdr:
    c.paragraphs[0].runs[0].bold = True

linhas_benef = [
    ("Gestão de múltiplas fontes de renda (o mais puro de 'contexto')",
     "Sintomas isolados não fecham diagnóstico. Plantão, consultório e sociedade também não — só vistos juntos.",
     "A IA vê pedaços soltos. A Caveo vê o corpo inteiro: 3 hospitais, consultório e sociedade num diagnóstico só."),
    ("Dashboard em tempo real",
     "Você monitora os sinais vitais do paciente ao vivo. E os seus, financeiros?",
     "Pergunta pra IA quanto você ganhou esse mês. Ela não sabe. Seu dashboard Caveo sabe — em tempo real."),
    ("Emissão ilimitada de NF em segundos",
     "Receita você assina em segundos. Nota fiscal devia ser igual.",
     "A IA te explica como emitir nota. A Caveo simplesmente emite — ilimitadas, em segundos."),
    ("Pagamento de tributos pelo app",
     "Tratamento sem adesão não funciona. Imposto esquecido também vira complicação.",
     "A IA te lembra que tem imposto a pagar. A Caveo paga por você, em 1 toque."),
    ("Economia de até 30% (o desfecho)",
     "Não é remédio novo — é o tratamento certo pro seu caso. Até 30% menos imposto.",
     "Resposta certa exige o caso certo. Com seu contexto completo, a Caveo corta até 30%."),
    ("600h/ano → 15 min/mês",
     "600h/ano de burocracia = ~25 plantões que você deixou de fazer. A Caveo te devolve esse tempo.",
     "Você não tem 600 horas pra explicar tudo pra uma IA. A Caveo já tem seus dados — resolve em 15 min/mês."),
    ("Suporte humano especializado",
     "No fim, quem entende de médico é médico — não um protocolo genérico.",
     "IA é ótima pra ser rápida. Mas o seu caso quem resolve é gente que entende plantão, não um bot."),
    ("Abertura de PJ grátis (recém-formado)",
     "Antes do primeiro plantão, a 'consulta inicial': abrir sua PJ. De graça.",
     "A IA te dá 12 passos. A Caveo abre sua PJ — com seus dados — em 1 botão."),
]
for a, b, c in linhas_benef:
    row = t.add_row().cells
    row[0].text = a
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = b
    row[2].text = c
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Padrão de campanha: ").bold = True
p.add_run(
    "um conceito-guarda-chuva com 8+ execuções, uma por benefício. Pode rodar como sequência de "
    "retargeting (dashboard → NF → múltiplas fontes → 30%) ou testar qual benefício puxa mais lead."
)

doc.add_page_break()

# ----------------------------------------------------------------------------
# 6. NOTAS DE EXECUÇÃO
# ----------------------------------------------------------------------------
doc.add_heading("6. Notas de execução", level=1)
bullet("Recém-formados convertem melhor com vídeo/depoimento (educar + criar urgência); maduros convertem com estático de comparação (\"Caveo vs. contador tradicional\") e número-choque (30%).")
bullet('Ofertas mais concretas para topo de funil em cada público: "30% de economia" (maduro) e "CNPJ grátis" (recém-formado).')
bullet("Segmentar recém-formados em: último ano/residente (ângulo antecipação) vs. recém-formado puro (ângulo faturar já).")
bullet('Os territórios mais ownáveis: "print de conversa com IA" (pattern interrupt) e "enquadramento clínico/contexto" (linguagem nativa + moat real).')

doc.save("/Users/matheus/Documents/Claude/Projects/caveo_analist_ai/docs/Ideias_Criativos_Anuncios_Caveo.docx")
print("OK: docs/Ideias_Criativos_Anuncios_Caveo.docx")
